"""
Flask Web Application for the Research Agent.

Serves the HTML/CSS frontend with authentication, chat history sidebar,
bridges to the backend agent orchestrator, and exposes a link-fetch API.
"""
import json
import os
import re
import io
import base64
from datetime import datetime

from flask import Flask, render_template, request, redirect, url_for, Response, send_file, flash, jsonify
from flask_login import login_required, login_user, logout_user, current_user
import markdown
import requests

from agent import run_research, continue_research, get_session
from auth import init_auth, authenticate, User
from config import Config
from link_fetcher import fetch_live_links
from database import (
    init_db,
    create_user,
    get_user_by_username,
    save_chat_history,
    list_chat_history,
    get_chat_history,
    get_chat_by_agent_session,
)

# matplotlib for chart generation
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

app = Flask(__name__)
app.secret_key = Config.SECRET_KEY
init_auth(app)
init_db()

# Make current_user always available in every Jinja2 template globally.
# This is the reliable fix for jinja2.exceptions.UndefinedError: 'current_user' is undefined.
app.jinja_env.globals['current_user'] = current_user


@app.context_processor
def inject_sidebar():
    if current_user.is_authenticated:
        try:
            return {"chat_history": list_chat_history(current_user.id)}
        except Exception:
            return {"chat_history": []}
    return {"chat_history": []}


def _make_citation_links_clickable(html: str) -> str:
    def replace(match):
        href = match.group(1)
        label = match.group(2)
        return (
            f'<a href="{href}" target="_blank" rel="noopener noreferrer" '
            f'class="citation-link" data-citation-id="{label}">{label}</a>'
        )
    return re.sub(r'<a href="([^\"]+)">(\[[0-9]+\])</a>', replace, html)


def _chart_to_base64(fig_plt) -> str:
    buf = io.BytesIO()
    fig_plt.savefig(buf, format='png', bbox_inches='tight', facecolor='none')
    buf.seek(0)
    img_b64 = base64.b64encode(buf.read()).decode('utf-8')
    plt.close()
    return img_b64


def generate_charts(data: dict) -> dict:
    charts = {'progress': '', 'reliability': '', 'coverage': ''}
    if Config.FAST_RESEARCH:
        return charts

    iterations = []
    facts_per_iter = []
    for entry in data.get('iteration_history', []):
        iterations.append(f"Iter {entry.get('iteration')}")
        facts_per_iter.append(entry.get('facts_extracted', 0))

    if iterations:
        fig, ax = plt.subplots(figsize=(6, 3))
        ax.bar(iterations, facts_per_iter, color='#6366F1')
        ax.set_title('Research Progress')
        ax.set_ylabel('Facts Extracted')
        charts['progress'] = _chart_to_base64(fig)

    source_labels = []
    source_scores = []
    for src in data.get('source_details', []):
        label = src.get('title') or src.get('url') or src.get('id', '')
        source_labels.append(label if len(label) <= 40 else label[:37] + '...')
        source_scores.append(int(src.get('score', 0)) * 10)

    if source_labels:
        fig, ax = plt.subplots(figsize=(6, max(2, len(source_labels) * 0.5)))
        ax.barh(source_labels, source_scores, color='#10B981')
        ax.set_title('Source Reliability')
        ax.set_xlabel('Reliability (0-100)')
        charts['reliability'] = _chart_to_base64(fig)

    return charts


def _enrich_results(results_data: dict) -> dict:
    md = markdown.Markdown(extensions=['fenced_code', 'tables'])
    report_source = results_data.get("report", "")
    if isinstance(report_source, dict):
        report_source = report_source.get("report", "")
    elif not isinstance(report_source, str):
        report_source = str(report_source)

    report_html = md.convert(report_source)
    results_data["report_html"] = _make_citation_links_clickable(report_html)

    if not Config.FAST_RESEARCH:
        try:
            prompt = f"Professional illustration of {results_data.get('question', 'research topic')}, digital art, clean"
            safe = requests.utils.requote_uri(prompt)
            results_data['image_url'] = f"https://image.pollinations.ai/prompt/{safe}?width=800&height=400&nologo=true"
        except Exception:
            results_data['image_url'] = None
    else:
        results_data['image_url'] = None

    try:
        results_data['charts'] = generate_charts(results_data)
    except Exception:
        results_data['charts'] = {'progress': '', 'reliability': '', 'coverage': ''}

    return results_data


def _resolve_export_data(session_id: str, user_id: int):
    """Returns (question, report_text, source_details, stats_dict) for export."""
    session = get_session(session_id)
    if session:
        return (
            session.current_question or "Research Report",
            session.report_payload.get("report", ""),
            session.source_details if hasattr(session, "source_details") else [],
            {
                "iterations": len(session.iteration_history) if hasattr(session, "iteration_history") else 0,
                "sources": len(session.sources) if hasattr(session, "sources") else 0,
                "facts": len(session.memory_snippets) if hasattr(session, "memory_snippets") else 0,
            },
        )

    saved = get_chat_by_agent_session(user_id, session_id)
    if saved:
        results = saved.get("results", {})
        return (
            saved.get("question") or "Research Report",
            saved.get("report_text") or results.get("report", ""),
            results.get("source_details", []),
            {
                "iterations": results.get("iterations", 0),
                "sources": results.get("sources_count", 0),
                "facts": results.get("facts_count", 0),
            },
        )
    return None


def _build_pdf(question: str, report_text: str, source_details: list, stats: dict) -> io.BytesIO:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab is not installed")

    now = datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(question, styles['Title']),
        Paragraph(f'Generated: {now}', styles['Normal']),
        Spacer(1, 12),
    ]

    tbl = Table(
        [['Iterations', str(stats.get('iterations', 0))],
         ['Sources', str(stats.get('sources', 0))],
         ['Facts', str(stats.get('facts', 0))]],
        hAlign='LEFT',
    )
    tbl.setStyle(TableStyle([('BACKGROUND', (0, 0), (1, 0), colors.whitesmoke), ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)]))
    story.extend([tbl, Spacer(1, 12)])

    for line in report_text.splitlines():
        story.append(Paragraph(line.replace('\t', '    '), styles['BodyText']))
    story.append(Spacer(1, 12))

    story.append(Paragraph('Sources', styles['Heading2']))
    for idx, src in enumerate(source_details, start=1):
        url = src.get('url') if isinstance(src, dict) else str(src)
        story.append(Paragraph(f"{idx}. {url}", styles['Normal']))

    story.append(Spacer(1, 24))
    story.append(Paragraph('Generated by ResearchAI', styles['Normal']))
    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route("/links", methods=["GET"])
@login_required
def links_page():
    return render_template("links.html", current_user=current_user)


@app.route("/api/links", methods=["GET", "POST"])
def api_fetch_links():
    """
    Extract and fetch live project links for a search query.

    GET  /api/links?q=your+query&max=5&validate=true
    POST JSON: {"query": "...", "max_results": 5, "validate": true}
    """
    if request.method == "POST":
        payload = request.get_json(silent=True) or {}
        query = (payload.get("query") or payload.get("q") or "").strip()
        max_results = payload.get("max_results", payload.get("max", Config.MAX_SEARCH_RESULTS))
        validate = payload.get("validate", True)
    else:
        query = (request.args.get("q") or request.args.get("query") or "").strip()
        max_results = request.args.get("max", Config.MAX_SEARCH_RESULTS, type=int)
        validate = request.args.get("validate", "true").lower() in ("1", "true", "yes")

    if not query:
        return jsonify({"error": "Missing query. Pass ?q=... or POST {\"query\": \"...\"}"}), 400

    try:
        links = fetch_live_links(query, max_results=max_results, validate=validate)
        return jsonify({
            "query": query,
            "count": len(links),
            "links": links,
        })
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/health")
def health():
    return jsonify({"status": "ok", "service": "ResearchAI"})


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        user = authenticate(username, password)
        if user:
            login_user(user)
            return redirect(url_for("index"))
        flash("Invalid username or password.", "error")

    return render_template("login.html", current_user=current_user)


@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return redirect(url_for("index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        confirm = request.form.get("confirm_password", "")

        if len(username) < 3:
            flash("Username must be at least 3 characters.", "error")
        elif len(password) < 6:
            flash("Password must be at least 6 characters.", "error")
        elif password != confirm:
            flash("Passwords do not match.", "error")
        elif get_user_by_username(username):
            flash("Username already taken.", "error")
        else:
            user_id = create_user(username, password)
            if user_id:
                login_user(User(user_id, username.lower()))
                flash("Account created successfully!", "success")
                return redirect(url_for("index"))
            flash("Could not create account.", "error")

    return render_template("register.html", current_user=current_user)


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("You have been logged out.", "success")
    return redirect(url_for("login"))


@app.route("/", methods=["GET"])
@login_required
def index():
    return render_template("index.html", current_user=current_user)


@app.route("/history/<int:history_id>", methods=["GET"])
@login_required
def view_history(history_id: int):
    entry = get_chat_history(current_user.id, history_id)
    if not entry:
        flash("Chat not found.", "error")
        return redirect(url_for("index"))

    results_data = entry.get("results") or {}
    if not results_data.get("report"):
        results_data["report"] = entry.get("report_text", "")
    results_data["session_id"] = entry.get("agent_session_id")
    results_data = _enrich_results(results_data)

    return render_template(
        "results.html",
        data=results_data,
        active_history_id=history_id,
        current_user=current_user
    )


@app.route("/research", methods=["POST"])
@login_required
def research():
    action = request.form.get("action")
    session_id = request.form.get("session_id")
    follow_up_prompt = request.form.get("follow_up_prompt")
    edited_question = request.form.get("edited_question")
    question = request.form.get("question")

    if action == "continue" and session_id:
        results_data = continue_research(session_id, follow_up_prompt=follow_up_prompt)
    elif action == "edit" and session_id and edited_question:
        results_data = continue_research(session_id, edited_question=edited_question)
    elif question:
        results_data = run_research(question)
    else:
        return redirect(url_for("index"))

    results_data = _enrich_results(results_data)

    if results_data.get("session_id") and results_data.get("question"):
        save_chat_history(
            current_user.id,
            results_data["session_id"],
            results_data["question"],
            results_data,
        )

    return render_template("results.html", data=results_data, current_user=current_user)


@app.route("/reresearch", methods=["POST"])
@login_required
def reresearch():
    session_id = request.form.get('session_id')
    edited_question = request.form.get('edited_question')
    focus_text = request.form.get('focus_text')

    if not session_id:
        return redirect(url_for('index'))

    if focus_text:
        edited_question = (edited_question or '') + ' Focus: ' + focus_text

    results_data = continue_research(session_id, edited_question=edited_question)
    results_data = _enrich_results(results_data)

    if results_data.get("session_id"):
        save_chat_history(
            current_user.id,
            results_data["session_id"],
            results_data.get("question", ""),
            results_data,
        )

    return render_template('results.html', data=results_data, current_user=current_user)


@app.route("/export/<string:export_format>", methods=["GET"])
@login_required
def export_report(export_format: str):
    session_id = request.args.get("session_id")
    if not session_id:
        return redirect(url_for("index"))

    resolved = _resolve_export_data(session_id, current_user.id)
    if not resolved:
        flash("Report not found. It may have expired — open it from your history sidebar.", "error")
        return redirect(url_for("index"))

    question, report_text, source_details, stats = resolved
    safe_title = re.sub(r'[^a-zA-Z0-9_-]', '_', question)[:40]

    if export_format == "markdown":
        return Response(
            report_text,
            mimetype="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=research_{safe_title}.md"},
        )

    if export_format == "docx":
        try:
            from docx import Document
        except ImportError:
            return Response("DOCX export requires: pip install python-docx", status=500)

        document = Document()
        document.add_heading(question, level=1)
        for line in report_text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                document.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.strip() == "":
                document.add_paragraph("")
            else:
                document.add_paragraph(line)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name=f"research_{safe_title}.docx",
            as_attachment=True,
        )

    if export_format == "pdf":
        try:
            buffer = _build_pdf(question, report_text, source_details, stats)
        except RuntimeError as e:
            return Response(str(e), status=500)
        return send_file(
            buffer,
            mimetype='application/pdf',
            download_name=f"research_{safe_title}.pdf",
            as_attachment=True,
        )

    if export_format == "txt":
        return Response(
            report_text,
            mimetype="text/plain",
            headers={"Content-Disposition": f"attachment; filename=research_{safe_title}.txt"},
        )

    flash("Unknown export format.", "error")
    return redirect(url_for("index"))


@app.route('/download/<string:session_id>', methods=['GET'])
@login_required
def download_pdf(session_id: str):
    resolved = _resolve_export_data(session_id, current_user.id)
    if not resolved:
        flash("Report not found.", "error")
        return redirect(url_for("index"))

    question, report_text, source_details, stats = resolved
    try:
        buffer = _build_pdf(question, report_text, source_details, stats)
    except RuntimeError as e:
        return Response(str(e), status=500)

    return send_file(
        buffer,
        mimetype='application/pdf',
        download_name=f"research_{session_id[:8]}.pdf",
        as_attachment=True,
    )


if __name__ == "__main__":
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "false").lower() in ("1", "true", "yes")

    print("\n" + "=" * 50)
    print("  ResearchAI is running!")
    print(f"  Local:  http://127.0.0.1:{port}")
    print(f"  Links:  http://127.0.0.1:{port}/api/links?q=your+topic")
    print("  Fast research mode:", Config.FAST_RESEARCH)
    print("  Press CTRL+C to stop the server")
    print("=" * 50 + "\n")
    app.run(host=host, port=port, debug=debug, use_reloader=debug)
